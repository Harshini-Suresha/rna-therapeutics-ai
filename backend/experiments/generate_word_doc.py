"""Generate a Word document summarizing all ASO platform experiments."""

import os
import json
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL

EXP_DIR = "backend/experiments"
OUTPUT_PATH = "docs/experiments_summary.docx"


def load_json(path):
    with open(path) as f:
        return json.load(f)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def main():
    doc = Document()

    doc.add_heading("ASO Platform - Experimental Results Summary", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("Dataset: OligoFormer/data/Hu.csv (2,361 ASO-mRNA pairs)")
    doc.add_paragraph("Seed: 42 | 5-fold cross-validation (KFold, shuffle=True)")
    doc.add_paragraph("Hyperparameters: lr=1e-3, dropout=0.2, weight_decay=1e-5, batch=32")
    doc.add_paragraph()

    # 1. Ablation Study
    doc.add_heading("1. Ablation Study", level=1)
    doc.add_paragraph("Which component of the model is actually responsible for the improvement?")

    ablation = load_json(os.path.join(EXP_DIR, "exp06_ablation", "full_results.json"))
    rows = []
    for exp in ablation:
        rows.append([
            exp["description"],
            f"{exp['pearson_mean']:.4f} +/- {exp['pearson_sd']:.4f}",
            f"{exp['r2_mean']:.4f} +/- {exp['r2_sd']:.4f}",
        ])
    add_table(doc, ["Feature Configuration", "Pearson", "R2"], rows)
    doc.add_paragraph(
        "Key finding: RNA-FM embeddings alone (Pearson=0.564) outperform "
        "the full FusionNet (Pearson=0.552). The plain MLP cannot effectively "
        "utilize the 20 handcrafted/accessibility features among 1300 total "
        "dimensions."
    )
    doc.add_paragraph()

    # 2. Baseline Comparison
    doc.add_heading("2. Baseline Comparison", level=1)
    doc.add_paragraph("FusionNet vs classical ML models on the full 1300-dim input.")

    baselines = load_json(os.path.join(EXP_DIR, "exp07_baselines", "metrics.json"))
    rows = []
    for name, res in baselines.items():
        label = name.replace("Baseline: ", "") if "Baseline" in name else name
        rows.append([
            label,
            f"{res['pearson_mean']:.4f} +/- {res['pearson_sd']:.4f}",
            f"{res['r2_mean']:.4f} +/- {res['r2_sd']:.4f}",
            f"{res['mse_mean']:.4f}",
            f"{res['mae_mean']:.4f}",
        ])
    add_table(doc, ["Model", "Pearson", "R2", "MSE", "MAE"], rows)
    doc.add_paragraph(
        "Key finding: FusionNet outperforms all classical baselines on Pearson "
        "correlation (0.552 vs 0.540 for SVR) with the lowest fold-to-fold "
        "variance (+/-0.018 vs +/-0.032 for SVR)."
    )
    doc.add_paragraph()

    # 3. SHAP
    doc.add_heading("3. SHAP Explainability", level=1)
    doc.add_paragraph("Why does FusionNet make its predictions?")

    shap_data = load_json(os.path.join(EXP_DIR, "exp08_shap", "family_importance.json"))
    rows = []
    for k, v in sorted(shap_data["normalized"].items(), key=lambda x: -x[1]):
        rows.append([k, f"{shap_data['raw'][k]:.4f}", f"{v*100:.1f}%"])
    add_table(doc, ["Feature Family", "Importance (sum |SHAP|)", "Normalized"], rows)
    doc.add_paragraph(
        "Key finding: RNA-FM embeddings account for ~99% of model explainability. "
        "The biological features (handcrafted + accessibility) are effectively "
        "ignored by the plain MLP - not because they're useless, but because the "
        "architecture cannot utilize them."
    )
    doc.add_paragraph()

    # 4. Feature Attention
    doc.add_heading("4. Feature Attention (Gated Fusion)", level=1)
    doc.add_paragraph("Can attention improve utilization of handcrafted biological descriptors?")

    fa = load_json(os.path.join(EXP_DIR, "exp09_feature_attention", "comparison.json"))
    delta_p = fa["gated"]["pearson_mean"] - fa["fusion"]["pearson_mean"]
    delta_r2 = fa["gated"]["r2_mean"] - fa["fusion"]["r2_mean"]

    rows = []
    for model_name in ["fusion", "gated"]:
        r = fa[model_name]
        label = "FusionNet" if model_name == "fusion" else "GatedFusionNet"
        delta_label = "-" if model_name == "fusion" else f"+{delta_p:.4f}"
        delta_r2_label = "-" if model_name == "fusion" else f"+{delta_r2:.4f}"
        rows.append([
            label,
            f"{r['pearson_mean']:.4f} +/- {r['pearson_sd']:.4f}",
            f"{r['r2_mean']:.4f} +/- {r['r2_sd']:.4f}",
            delta_label,
            delta_r2_label,
        ])
    add_table(doc, ["Model", "Pearson", "R2", "DeltaPearson", "DeltaR2"], rows)

    doc.add_paragraph()
    doc.add_paragraph("Learned per-family attention (gate) weights:")

    gates = fa["gated"]["avg_gate_weights"]
    total = sum(gates.values())
    gate_rows = []
    for fam_name, val in sorted(gates.items(), key=lambda x: -x[1]):
        gate_rows.append([fam_name, f"{val:.4f}", f"{val/total*100:.1f}%"])
    add_table(doc, ["Feature Family", "Gate Weight", "Normalized"], gate_rows)

    doc.add_paragraph(
        f"Key finding: GatedFusionNet improves Pearson by +{delta_p:.4f} "
        f"(0.552 -> 0.585) and R2 by +{delta_r2:.4f} "
        f"(0.274 -> 0.331). The attention mechanism gives meaningful weights "
        f"to all 4 feature families - accessibility and handcrafted features "
        f"receive ~20-22% attention, versus <1% SHAP importance in the plain MLP. "
        "This confirms the biological features contain useful signal that the "
        "plain concatenation + MLP was failing to leverage."
    )
    doc.add_paragraph()

    # 5. Unified Results
    doc.add_heading("5. Unified Results", level=1)
    all_rows = []
    for key in ["exp01_handcrafted", "exp02_rnafm", "exp03_accessibility",
                "exp04_rnafm_accessibility", "exp05_rnafm_handcrafted", "exp06_fusion"]:
        exp = next(e for e in ablation if e["experiment"] == key)
        all_rows.append([
            exp["description"][:30],
            "Ablation",
            f"{exp['pearson_mean']:.4f} +/- {exp['pearson_sd']:.4f}",
            f"{exp['r2_mean']:.4f} +/- {exp['r2_sd']:.4f}",
        ])

    for name, res in baselines.items():
        label = name.replace("Baseline: ", "") if "Baseline" in name else name
        all_rows.append([
            label[:30],
            "Baseline",
            f"{res['pearson_mean']:.4f} +/- {res['pearson_sd']:.4f}",
            f"{res['r2_mean']:.4f} +/- {res['r2_sd']:.4f}",
        ])

    all_rows.append([
        "FusionNet (attention)",
        "FeatureAttention",
        f"{fa['fusion']['pearson_mean']:.4f} +/- {fa['fusion']['pearson_sd']:.4f}",
        f"{fa['fusion']['r2_mean']:.4f} +/- {fa['fusion']['r2_sd']:.4f}",
    ])

    all_rows.append([
        "GatedFusionNet (BEST)",
        "FeatureAttention",
        f"{fa['gated']['pearson_mean']:.4f} +/- {fa['gated']['pearson_sd']:.4f}",
        f"{fa['gated']['r2_mean']:.4f} +/- {fa['gated']['r2_sd']:.4f}",
    ])

    add_table(doc, ["Model", "Category", "Pearson", "R2"], all_rows)

    doc.add_paragraph()
    doc.add_paragraph(f"Current best result: Pearson = {fa['gated']['pearson_mean']:.4f} +/- {fa['gated']['pearson_sd']:.4f}")
    doc.add_paragraph("Original benchmark:  Pearson = 0.560 +/- 0.026 (FusionNet, HPO)")

    # 6. Next Steps
    doc.add_heading("6. Next Steps", level=1)
    doc.add_paragraph(
        "The scientific story is now well-motivated:"
    )
    doc.add_paragraph("  1. Ablation: RNA-FM is the dominant feature (0.564 vs 0.297 for handcrafted)")
    doc.add_paragraph("  2. SHAP: Biological features get <1% importance in plain MLP")
    doc.add_paragraph("  3. Feature Attention: GatedFusionNet boosts Pearson by +0.033")
    doc.add_paragraph("  4. Next: Cross-Attention Transformer to model guide-target interactions")

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Word document saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
