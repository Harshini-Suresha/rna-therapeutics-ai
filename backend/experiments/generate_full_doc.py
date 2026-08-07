"""Generate a comprehensive Word document containing ALL ML experiment observations."""

import os
import json
import csv
from datetime import datetime

import numpy as np
import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt

EXP_DIR = "backend/experiments"
RESULTS_DIR = "backend/results"
OUTPUT_PATH = "experiments_observations_full.docx"


def load_json(path):
    with open(path) as f:
        return json.load(f)


def add_table(doc, headers, rows, col_widths=None, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(font_size)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(font_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def add_picture(doc, path, width=5.5):
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(width))
        except Exception as e:
            doc.add_paragraph(f"[Could not embed image: {path}] ({e})")


def fold_rows(metrics):
    rows = []
    for f in metrics.get("folds", []):
        rows.append([
            f["fold"],
            f"{f['pearson']:.4f}",
            f"{f['r2']:.4f}",
            f"{f['mse']:.4f}",
            f"{f['mae']:.4f}",
        ])
    return rows


def per_experiment_section(doc, exp_key, label):
    exp_dir = os.path.join(EXP_DIR, exp_key)
    m = load_json(os.path.join(exp_dir, "metrics.json"))
    doc.add_heading(f"{label} ({exp_key})", level=3)
    doc.add_paragraph(
        f"Description: {m['description']} | Feature families: {', '.join(m['feature_families'])} "
        f"| Input dim: {m['input_dim']} | Architecture: {m['architecture']}"
    )
    add_table(doc, ["Fold", "Pearson", "R²", "MSE", "MAE"], fold_rows(m))
    doc.add_paragraph(
        f"Summary: Pearson = {m['pearson_mean']:.4f} ± {m['pearson_sd']:.4f} | "
        f"R² = {m['r2_mean']:.4f} ± {m['r2_sd']:.4f} | "
        f"MSE = {m['mse_mean']:.4f} ± {m['mse_sd']:.4f} | "
        f"MAE = {m['mae_mean']:.4f} ± {m['mae_sd']:.4f}"
    )
    curve = os.path.join(exp_dir, "training_curve.png")
    if os.path.exists(curve):
        doc.add_paragraph("Training curve:")
        add_picture(doc, curve, width=5.0)
    cfg = os.path.join(exp_dir, "config.yaml")
    if os.path.exists(cfg):
        doc.add_paragraph("Configuration:")
        with open(cfg) as f:
            doc.add_paragraph(f.read(), style="No Spacing")


def main():
    doc = Document()

    doc.add_heading("ASO Platform — Complete ML Observations & Experimental Results", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("This document contains every recorded observation, metric, and "
                      "configuration from the ASO efficacy prediction ML pipeline.")
    doc.add_paragraph()

    # ---------- 1. Overview / setup ----------
    doc.add_heading("1. Overview & Experimental Setup", level=1)
    doc.add_paragraph("Dataset: OligoFormer/data/Hu.csv (2,361 ASO-mRNA pairs)")
    doc.add_paragraph("Embedding cache: backend/data/hu_embeddings.pt (precomputed RNA-FM)")
    doc.add_paragraph("Seed: 42 | 5-fold cross-validation (KFold, shuffle=True)")
    doc.add_paragraph("Base hyperparameters (from HPO): lr=1e-3, dropout=0.2, weight_decay=1e-5, batch=32")
    doc.add_paragraph("Training: Adam optimizer, MSE loss, early stopping (patience=10), max 100 epochs")
    doc.add_paragraph()

    # ---------- 2. Model architectures ----------
    doc.add_heading("2. Model Architectures", level=1)
    doc.add_heading("FusionNet", level=2)
    doc.add_paragraph(
        "Plain concatenation + MLP. All feature families are concatenated into a single "
        "1300-dim vector (640-dim siRNA RNA-FM + 640-dim mRNA RNA-FM + 11 accessibility "
        "+ 9 handcrafted features) and passed through a stack of linear layers: "
        "1300 → 512 → 256 → 128 → 64 → 1. Each block is Linear + BatchNorm1d + GELU + Dropout."
    )
    doc.add_heading("GatedFusionNet", level=2)
    doc.add_paragraph(
        "Replaces concatenation with learned per-family gating. Each family is projected "
        "to 128 dims, a sigmoid gate weight is learned per sample per family, gated features "
        "are concatenated (4×128) and fed through the same MLP (512→256→128→64→1). "
        "Family dims: [640 siRNA, 640 mRNA, 11 accessibility, 9 handcrafted]."
    )
    doc.add_paragraph()

    # ---------- 3. Hyperparameter search ----------
    doc.add_heading("3. Hyperparameter Search (HPO)", level=1)
    doc.add_paragraph(
        "Grid search over 243 configurations (lr ∈ {1e-3, 5e-4, 1e-4}, dropout ∈ {0.2, 0.3, 0.4}, "
        "hidden ∈ {[256,128,64], [512,256,128,64], [768,512,256]}, weight_decay ∈ {0, 1e-5, 1e-4}, "
        "batch ∈ {16, 32, 64}), evaluated with 3-fold CV, 30 epochs/fold, patience 10."
    )
    hpo_path = os.path.join(RESULTS_DIR, "hyperparameter_results.csv")
    rows = []
    with open(hpo_path) as f:
        reader = csv.DictReader(f)
        for i, r in enumerate(reader, 1):
            rows.append([
                i,
                r["LR"], r["Dropout"], r["Hidden"], r["WD"], r["Batch"],
                f"{float(r['Pearson_Mean']):.4f}", f"{float(r['Pearson_SD']):.4f}",
                f"{float(r['R2_Mean']):.4f}", f"{float(r['R2_SD']):.4f}",
                r["Time_s"],
            ])
    add_table(doc, ["Run", "LR", "Dropout", "Hidden", "WD", "Batch",
                    "Pearson", "SD", "R²", "SD", "Time(s)"], rows,
              col_widths=[0.4, 0.5, 0.5, 1.1, 0.6, 0.5, 0.7, 0.6, 0.7, 0.6, 0.6])
    best = max(rows, key=lambda r: float(r[6]))
    doc.add_paragraph(
        f"Best HPO config (run {best[0]}): lr={best[1]}, dropout={best[2]}, "
        f"hidden={best[3]}, weight_decay={best[4]}, batch={best[5]} "
        f"→ Pearson = {best[6]} ± {best[7]}, R² = {best[8]} ± {best[9]}."
    )
    doc.add_paragraph(
        "Observation: the best config found was lr=0.001, dropout=0.2, "
        "hidden=[512, 256, 128, 64], weight_decay=1e-05, batch=16. "
        "batch=64 consistently underperformed across all configurations; "
        "larger hidden widths were not better."
    )
    doc.add_paragraph()

    # ---------- 4. Final model CV (Huesken) ----------
    doc.add_heading("4. Final FusionNet Cross-Validation (backend/results)", level=1)
    hm = load_json(os.path.join(RESULTS_DIR, "Huesken", "metrics.json"))
    doc.add_paragraph(
        f"Model: {hm['model']} | Dataset: {hm['dataset']} ({hm['samples']} samples) | "
        f"{hm['folds']} folds"
    )
    doc.add_paragraph(
        f"Pearson = {hm['pearson']} ± {hm['pearson_sd']} | R² = {hm['r2']} ± {hm['r2_sd']} | "
        f"MSE = {hm['mse']} | MAE = {hm['mae']} | epochs = {hm['epochs']}"
    )
    doc.add_paragraph(
        f"Best fold: {hm['best_fold']} (Pearson {hm['best_fold_pearson']}) | "
        f"OOF: Pearson {hm['oof_pearson']}, R² {hm['oof_r2']}, "
        f"MSE {hm['oof_mse']}, MAE {hm['oof_mae']}"
    )
    cv_rows = []
    with open(os.path.join(RESULTS_DIR, "cv_results.csv")) as f:
        for line in f:
            cv_rows.append(line.strip().split(","))
    add_table(doc, cv_rows[0], cv_rows[1:], col_widths=[1.0, 1.0, 1.0, 1.0, 1.0, 1.0])
    doc.add_paragraph()
    doc.add_heading("Training log (final model)", level=2)
    log_rows = []
    with open(os.path.join(RESULTS_DIR, "Huesken", "training_log.csv")) as f:
        reader = csv.DictReader(f)
        for r in reader:
            log_rows.append([r["epoch"], f"{float(r['train_loss']):.4f}",
                             f"{float(r['val_loss']):.4f}",
                             f"{float(r['train_pearson']):.4f}",
                             f"{float(r['val_pearson']):.4f}"])
    add_table(doc, ["Epoch", "TrainLoss", "ValLoss", "TrainPearson", "ValPearson"], log_rows)
    doc.add_paragraph("Learning curve:")
    add_picture(doc, os.path.join(RESULTS_DIR, "Huesken", "learning_curve.png"), width=5.0)
    doc.add_paragraph("Predicted vs true (OOF) scatter:")
    add_picture(doc, os.path.join(RESULTS_DIR, "Huesken", "scatter_plot.png"), width=5.0)
    doc.add_paragraph()
    doc.add_heading("Binary classification evaluation (OOF, threshold 0.5)", level=2)
    doc.add_paragraph(open(os.path.join(RESULTS_DIR, "Huesken", "confusion.txt")).read(),
                      style="No Spacing")
    doc.add_paragraph("Per-sample OOF predictions (Sample, True, Predicted) are stored in "
                      "backend/results/Huesken/predictions.csv (2,361 rows).")
    doc.add_paragraph()

    # ---------- 5. CV analysis ----------
    doc.add_heading("5. CV Analysis & Training Improvements", level=1)
    doc.add_paragraph(open(os.path.join(EXP_DIR, "cv_analysis.md")).read(), style="No Spacing")
    doc.add_paragraph("Raw CV results:")
    doc.add_paragraph(open(os.path.join(EXP_DIR, "cv_results.txt")).read(), style="No Spacing")
    doc.add_paragraph()

    # ---------- 6. Ablation study ----------
    doc.add_heading("6. Ablation Study", level=1)
    ablation = load_json(os.path.join(EXP_DIR, "exp06_ablation", "full_results.json"))
    summary_rows = []
    for exp in ablation:
        summary_rows.append([
            exp["description"],
            f"{exp['pearson_mean']:.4f} ± {exp['pearson_sd']:.4f}",
            f"{exp['r2_mean']:.4f} ± {exp['r2_sd']:.4f}",
        ])
    add_table(doc, ["Feature Configuration", "Pearson", "R²"], summary_rows)
    doc.add_paragraph(
        "Key finding: RNA-FM embeddings alone (Pearson=0.564) outperform the full FusionNet "
        "(0.552). Handcrafted features contribute 0.297, accessibility only 0.119. "
        "Concatenation hurts: RNA-FM + anything ≤ RNA-FM alone."
    )
    for exp in ablation:
        per_experiment_section(doc, exp["experiment"], exp["description"])
    doc.add_paragraph()

    # ---------- 7. Baselines ----------
    doc.add_heading("7. Baseline Comparison", level=1)
    baselines = load_json(os.path.join(EXP_DIR, "exp07_baselines", "metrics.json"))
    summary_rows = []
    for name, res in baselines.items():
        summary_rows.append([
            name,
            f"{res['pearson_mean']:.4f} ± {res['pearson_sd']:.4f}",
            f"{res['r2_mean']:.4f} ± {res['r2_sd']:.4f}",
            f"{res['mse_mean']:.4f}",
            f"{res['mae_mean']:.4f}",
        ])
    add_table(doc, ["Model", "Pearson", "R²", "MSE", "MAE"], summary_rows)
    doc.add_paragraph(
        "Key finding: FusionNet edges SVR on Pearson (0.552 vs 0.540) with much lower "
        "fold-to-fold variance (±0.018 vs ±0.032). RandomForest and XGBoost are clearly worse."
    )
    for name, res in baselines.items():
        doc.add_heading(f"Per-fold details — {name}", level=3)
        add_table(doc, ["Fold", "Pearson", "R²", "MSE", "MAE"], fold_rows(res))
    doc.add_paragraph()

    # ---------- 8. SHAP ----------
    doc.add_heading("8. SHAP Explainability", level=1)
    shap = load_json(os.path.join(EXP_DIR, "exp08_shap", "metrics.json"))
    doc.add_paragraph(
        f"Explainer: {shap['explainer']} | Surrogate: {shap['surrogate_model']} "
        f"(R²={shap['surrogate_r2']:.4f}) | Background samples: {shap['n_background']} | "
        f"Explained samples: {shap['n_explain']}"
    )
    doc.add_heading("Feature family importance", level=2)
    fam_rows = []
    for k, v in sorted(shap["family_importance_normalized"].items(), key=lambda x: -x[1]):
        fam_rows.append([k, f"{shap['family_importance_raw'][k]:.4f}", f"{v*100:.2f}%"])
    add_table(doc, ["Feature Family", "Importance (sum |SHAP|)", "Normalized"], fam_rows)
    doc.add_paragraph(
        "Key finding: RNA-FM embeddings account for ~99% of model explainability "
        "(mRNA 50.5% + siRNA 48.7%). Handcrafted (0.5%) and accessibility (0.3%) "
        "are effectively ignored by the plain MLP."
    )
    doc.add_heading("Top 20 individual features", level=2)
    top_rows = [[t["feature"], f"{t['mean_abs_shap']:.4f}"] for t in shap["top_features"]]
    add_table(doc, ["Feature", "Mean |SHAP|"], top_rows)
    doc.add_heading("Prediction statistics (explained sample)", level=2)
    ps = shap["prediction_stats"]
    add_table(doc, ["", "Mean", "Std"],
              [["Predicted", f"{ps['pred_mean']:.4f}", f"{ps['pred_std']:.4f}"],
               ["True", f"{ps['true_mean']:.4f}", f"{ps['true_std']:.4f}"]])
    doc.add_paragraph("SHAP plots:")
    for img in ["shap_summary_by_family.png", "shap_family_importance.png", "shap_feature_importance.png"]:
        add_picture(doc, os.path.join(EXP_DIR, "exp08_shap", img), width=5.0)
    doc.add_heading("Force plots", level=2)
    fp_dir = os.path.join(EXP_DIR, "exp08_shap", "force_plots")
    if os.path.isdir(fp_dir):
        for img in sorted(os.listdir(fp_dir)):
            if img.endswith(".png"):
                doc.add_paragraph(img)
                add_picture(doc, os.path.join(fp_dir, img), width=5.5)
    doc.add_paragraph()

    # ---------- 9. Feature attention ----------
    doc.add_heading("9. Feature Attention (Gated Fusion)", level=1)
    fa = load_json(os.path.join(EXP_DIR, "exp09_feature_attention", "comparison.json"))
    delta_p = fa["gated"]["pearson_mean"] - fa["fusion"]["pearson_mean"]
    delta_r2 = fa["gated"]["r2_mean"] - fa["fusion"]["r2_mean"]
    rows = []
    for model_name in ["fusion", "gated"]:
        r = fa[model_name]
        label = "FusionNet" if model_name == "fusion" else "GatedFusionNet"
        rows.append([
            label,
            f"{r['pearson_mean']:.4f} ± {r['pearson_sd']:.4f}",
            f"{r['r2_mean']:.4f} ± {r['r2_sd']:.4f}",
            "-" if model_name == "fusion" else f"+{delta_p:.4f}",
            "-" if model_name == "fusion" else f"+{delta_r2:.4f}",
        ])
    add_table(doc, ["Model", "Pearson", "R²", "ΔPearson", "ΔR²"], rows)
    doc.add_paragraph(
        f"Key finding: GatedFusionNet improves Pearson by +{delta_p:.4f} (0.552 → 0.585) "
        f"and R² by +{delta_r2:.4f} (0.274 → 0.331). Learned gating lets the model "
        f"utilize biological features that concatenation ignored."
    )
    doc.add_heading("Learned gate weights", level=2)
    gates = fa["gated"]["avg_gate_weights"]
    total = sum(gates.values())
    gate_rows = []
    for fam, val in sorted(gates.items(), key=lambda x: -x[1]):
        gate_rows.append([fam, f"{val:.4f}", f"{val/total*100:.2f}%"])
    add_table(doc, ["Feature Family", "Gate Weight", "Normalized"], gate_rows)
    doc.add_paragraph(
        "Key finding: gate weights now spread across all families — accessibility 21.6% and "
        "handcrafted 19.5% — vs <1% SHAP importance in the plain MLP. The biological features "
        "contain useful signal the plain architecture was failing to leverage."
    )
    for model_name in ["fusion", "gated"]:
        m = load_json(os.path.join(EXP_DIR, "exp09_feature_attention", model_name, "metrics.json"))
        label = "FusionNet" if model_name == "fusion" else "GatedFusionNet"
        doc.add_heading(f"Per-fold details — {label}", level=3)
        add_table(doc, ["Fold", "Pearson", "R²", "MSE", "MAE"], fold_rows(m))
        curve = os.path.join(EXP_DIR, "exp09_feature_attention", model_name, "training_curve.png")
        if os.path.exists(curve):
            add_picture(doc, curve, width=5.0)
    doc.add_paragraph()

    # ---------- 9b. Cross-feature interaction ----------
    doc.add_heading("10. Cross-Feature Interaction (CrossAttentionFusion)", level=1)
    doc.add_paragraph(
        "Research question: can we explicitly model the interaction between the ASO and its "
        "target RNA instead of treating them as independent embeddings?"
    )
    doc.add_heading("Embedding check (design decision)", level=2)
    doc.add_paragraph(
        "The RNA-FM embeddings used in this pipeline are MEAN-POOLED: "
        "backend/features/rnafm.py runs token_embeddings.mean(dim=1), producing one fixed "
        "640-dim vector per sequence (stored in backend/data/hu_embeddings.pt as "
        "siRNA(640) + mRNA(640)). There are NO per-nucleotide token embeddings (L x 640)."
    )
    doc.add_paragraph(
        "Consequence: a true nucleotide-level cross-attention transformer is impossible with "
        "the cached features (nn.MultiheadAttention over a single token is a no-op). "
        "Per the prescribed fallback, we implemented a cross-feature interaction module "
        "instead: feature-level cross-attention (query = ASO, key/value = target) plus an "
        "elementwise bilinear interaction, plus gated accessibility/handcrafted features."
    )
    ca = load_json(os.path.join(EXP_DIR, "exp10_cross_attention", "comparison.json"))
    ca_metrics = load_json(os.path.join(EXP_DIR, "exp10_cross_attention", "metrics.json"))
    rows = []
    for name in ["rnafm_only", "fusion", "gated", "cross_attention"]:
        r = ca[name]
        label = {"rnafm_only": "RNA-FM only", "fusion": "FusionNet",
                 "gated": "GatedFusionNet", "cross_attention": "CrossAttentionFusion"}[name]
        rows.append([
            label,
            f"{r['pearson_mean']:.4f} ± {r['pearson_sd']:.4f}",
            f"{r['r2_mean']:.4f} ± {r['r2_sd']:.4f}",
        ])
    add_table(doc, ["Model", "Pearson", "R²"], rows)
    delta_ca = ca["cross_attention"]["pearson_mean"] - ca["gated"]["pearson_mean"]
    doc.add_paragraph(
        f"Result: CrossAttentionFusion reaches Pearson {ca['cross_attention']['pearson_mean']:.4f} "
        f"± {ca['cross_attention']['pearson_sd']:.4f}, statistically indistinguishable from "
        f"GatedFusionNet ({delta_ca:+.4f} ΔPearson) and clearly above FusionNet and RNA-FM only. "
        "The interaction terms do not add signal beyond learned gating when embeddings are "
        "already pooled to a single vector per sequence — consistent with the hypothesis that "
        "nucleotide-level interaction is where the remaining signal would be."
    )
    doc.add_heading("Per-fold details", level=3)
    add_table(doc, ["Fold", "Pearson", "R²", "MSE", "MAE"], fold_rows(ca_metrics))
    doc.add_paragraph("Architecture: rnafm_dim=640, proj_dim=256, gate_hidden=32, "
                      "head=[512, 256, 128, 64]; input 1300-dim. Same hyperparameters as "
                      "all prior experiments (lr=1e-3, dropout=0.2, wd=1e-5, batch=32, "
                      "patience=10, max 100 epochs).")
    doc.add_paragraph("Training curve:")
    add_picture(doc, os.path.join(EXP_DIR, "exp10_cross_attention", "training_curve.png"), width=5.0)
    doc.add_paragraph("Learned cross-feature attention map (256x256, ASO features → target features):")
    add_picture(doc, os.path.join(EXP_DIR, "exp10_cross_attention", "cross_attention_heatmap.png"), width=5.0)
    doc.add_paragraph(
        "Limitation: this heatmap is feature-level (256 projected dims), not nucleotide-level. "
        "A biologically interpretable ASO-position × target-position map requires re-extracting "
        "token embeddings from RNA-FM (skip the mean pool) and feeding token sequences — "
        "left as the next phase."
    )
    doc.add_paragraph()

    # ---------- 11. Token-level cross-attention ----------
    doc.add_heading("11. Token-Level Cross-Attention Transformer (TokenCrossAttention)", level=1)
    doc.add_paragraph(
        "Research question: does modeling the ASO-target interaction at the NUCLEOTIDE level "
        "(instead of the single-vector, feature-level attention of Phase 10) improve prediction?"
    )
    doc.add_heading("Token representation (design decision)", level=2)
    doc.add_paragraph(
        "RNA-FM was re-run without the mean pool (backend/features/rnafm.py now exposes "
        "embed_with_tokens / embed_batch_with_tokens). Per-nucleotide embeddings are cached in "
        "backend/data/hu_token_embeddings.pt: aso_tokens (2361, 19, 640), mrna_tokens "
        "(2361, 57, 640)."
    )
    doc.add_paragraph(
        "Target representation: Huesken mRNA targets are already pre-windowed — the binding "
        "site is exactly the central 19 nt (verified: starts at position 19 of 57 in every "
        "sample), i.e. 19 nt upstream + 19 nt binding site + 19 nt downstream. The full "
        "57-nt window is used for training (fixed length -> no padding, no attention mask). "
        "For inference on full-length mRNAs (>1000 nt), a windowing step (e.g. ±40 nt around "
        "the predicted binding site) is applied at inference time only."
    )
    doc.add_paragraph(
        "Architecture: per-token projection to d_model=128 + learnable positional encoding "
        "(required: RNA-FM token embeddings carry no position info) -> target self-attention "
        "(1 transformer block, 4 heads) -> ASO-position cross-attention over target positions "
        "-> mean-pool + elementwise bilinear + gated accessibility/handcrafted features -> "
        "MLP head [512, 256, 128, 64]."
    )
    tc = load_json(os.path.join(EXP_DIR, "exp11_token_cross_attention", "comparison.json"))
    tc_metrics = load_json(os.path.join(EXP_DIR, "exp11_token_cross_attention", "metrics.json"))
    rows = []
    for name in ["rnafm_only", "fusion", "gated", "cross_attention", "token_cross_attention"]:
        r = tc[name]
        label = {"rnafm_only": "RNA-FM only", "fusion": "FusionNet",
                 "gated": "GatedFusionNet", "cross_attention": "CrossAttentionFusion (feat-level)",
                 "token_cross_attention": "TokenCrossAttention (nt-level)"}[name]
        rows.append([
            label,
            f"{r['pearson_mean']:.4f} ± {r['pearson_sd']:.4f}",
            f"{r['r2_mean']:.4f} ± {r['r2_sd']:.4f}",
        ])
    add_table(doc, ["Model", "Pearson", "R²"], rows)
    delta_tc = tc["token_cross_attention"]["pearson_mean"] - tc["gated"]["pearson_mean"]
    doc.add_paragraph(
        f"Result: TokenCrossAttention reaches Pearson {tc['token_cross_attention']['pearson_mean']:.4f} "
        f"± {tc['token_cross_attention']['pearson_sd']:.4f} and R² "
        f"{tc['token_cross_attention']['r2_mean']:.4f} ± {tc['token_cross_attention']['r2_sd']:.4f} — "
        f"the best result so far, {delta_tc:+.4f} ΔPearson over GatedFusionNet and +0.035 over "
        "CrossAttentionFusion. Per-fold paired t-test vs exp10: p=0.067 (trending, not significant "
        "at alpha=0.05 with 5 folds; direction consistent in 4 of 5 folds)."
    )
    doc.add_heading("Per-fold details", level=3)
    add_table(doc, ["Fold", "Pearson", "R²", "MSE", "MAE"], fold_rows(tc_metrics))
    doc.add_paragraph("Architecture: rnafm_dim=640, d_model=128, n_heads=4, aso_len=19, "
                      "target_len=57, gate_hidden=32, head=[512, 256, 128, 64]. Same "
                      "hyperparameters as all prior experiments (lr=1e-3, dropout=0.2, wd=1e-5, "
                      "batch=32, patience=10, max 100 epochs).")
    doc.add_paragraph("Training curve:")
    add_picture(doc, os.path.join(EXP_DIR, "exp11_token_cross_attention", "training_curve.png"), width=5.0)
    doc.add_paragraph(
        "ASO-position x target-position cross-attention matrix (averaged over heads, batches, "
        "and the 5 folds; red box = binding site, target positions 19-37):"
    )
    add_picture(doc, os.path.join(EXP_DIR, "exp11_token_cross_attention", "aso_x_target_attention.png"), width=5.0)
    attn = np.array(tc_metrics["avg_aso_x_target_attention"])
    doc.add_paragraph(
        f"Interpretation: the binding-site region receives the highest mean attention "
        f"({attn[:, 19:38].mean():.4f}) vs upstream ({attn[:, :19].mean():.4f}) and downstream "
        f"({attn[:, 38:].mean():.4f}) — the model learns to focus ASO positions on the "
        "complementary target region, a biologically sensible interaction map."
    )
    doc.add_paragraph()

    # ---------- 12. Unified table ----------
    doc.add_heading("12. Unified Results (all models)", level=1)
    all_rows = []
    for exp in ablation:
        all_rows.append([
            exp["description"],
            "Ablation",
            f"{exp['pearson_mean']:.4f} ± {exp['pearson_sd']:.4f}",
            f"{exp['r2_mean']:.4f} ± {exp['r2_sd']:.4f}",
        ])
    for name, res in baselines.items():
        all_rows.append([
            name,
            "Baseline",
            f"{res['pearson_mean']:.4f} ± {res['pearson_sd']:.4f}",
            f"{res['r2_mean']:.4f} ± {res['r2_sd']:.4f}",
        ])
    all_rows.append([
        "FusionNet (attention run)",
        "FeatureAttention",
        f"{fa['fusion']['pearson_mean']:.4f} ± {fa['fusion']['pearson_sd']:.4f}",
        f"{fa['fusion']['r2_mean']:.4f} ± {fa['fusion']['r2_sd']:.4f}",
    ])
    all_rows.append([
        "GatedFusionNet (BEST)",
        "FeatureAttention",
        f"{fa['gated']['pearson_mean']:.4f} ± {fa['gated']['pearson_sd']:.4f}",
        f"{fa['gated']['r2_mean']:.4f} ± {fa['gated']['r2_sd']:.4f}",
    ])
    all_rows.append([
        "CrossAttentionFusion",
        "CrossFeature",
        f"{ca['cross_attention']['pearson_mean']:.4f} ± {ca['cross_attention']['pearson_sd']:.4f}",
        f"{ca['cross_attention']['r2_mean']:.4f} ± {ca['cross_attention']['r2_sd']:.4f}",
    ])
    all_rows.append([
        "TokenCrossAttention (BEST)",
        "TokenCross",
        f"{tc['token_cross_attention']['pearson_mean']:.4f} ± {tc['token_cross_attention']['pearson_sd']:.4f}",
        f"{tc['token_cross_attention']['r2_mean']:.4f} ± {tc['token_cross_attention']['r2_sd']:.4f}",
    ])
    add_table(doc, ["Model", "Category", "Pearson", "R²"], all_rows)
    best_pearson = max(
        [fa["gated"]["pearson_mean"], ca["cross_attention"]["pearson_mean"],
         tc["token_cross_attention"]["pearson_mean"]]
    )
    doc.add_paragraph(
        f"Current best result: Pearson = {best_pearson:.4f} "
        f"(TokenCrossAttention {tc['token_cross_attention']['pearson_mean']:.4f} vs "
        f"GatedFusionNet {fa['gated']['pearson_mean']:.4f} vs "
        f"CrossAttentionFusion {ca['cross_attention']['pearson_mean']:.4f})"
    )
    doc.add_paragraph("Original benchmark: Pearson = 0.560 ± 0.026 (FusionNet, HPO)")
    doc.add_paragraph()

    # ---------- 13. Observations / conclusions ----------
    doc.add_heading("13. Summary of Observations", level=1)
    for obs in [
        "RNA-FM embeddings are the dominant predictive signal (Pearson 0.564 alone).",
        "Handcrafted features alone reach 0.297; accessibility alone only 0.119.",
        "Plain concatenation of RNA-FM with biological features does NOT help (0.552 ≤ 0.564).",
        "Classical ML: SVR approaches FusionNet (0.540 vs 0.552); RandomForest/XGBoost are far behind.",
        "SHAP shows biological features get <1% of plain-MLP importance.",
        "Gated attention recovers the biological signal: +0.033 Pearson, +0.057 R² over plain FusionNet.",
        "RNA-FM embeddings are mean-pooled to one 640-dim vector per sequence; no per-nucleotide tokens are cached.",
        "Feature-level cross-attention + bilinear interaction (CrossAttentionFusion) reaches 0.584 Pearson — statistically tied with GatedFusionNet, above FusionNet and RNA-FM-only.",
        "Interaction terms add no signal beyond gating when embeddings are pooled; remaining signal likely lives at nucleotide level.",
        "Token-level cross-attention (TokenCrossAttention) reaches Pearson 0.6195 ± 0.0264 — best result, +0.035 over GatedFusionNet and CrossAttentionFusion (paired t-test p=0.067, trending).",
        "Attention matrix shows the model focuses ASO positions on the binding-site region of the target (mean attention 0.0195 vs 0.0168/0.0164 up/downstream).",
        "batch=64 consistently underperforms; best HPO config uses batch=16/32.",
        "10-epoch training was insufficient; 100 epochs + early stopping raised every fold into 0.51–0.59.",
        "Final OOF model: Pearson 0.543, R² 0.289; binary threshold 0.5 gives 70.5% accuracy.",
    ]:
        doc.add_paragraph(f"• {obs}")

    doc.add_heading("14. Next Steps", level=1)
    doc.add_paragraph("  1. Increase statistical power (nested CV / repeated folds / bootstrap) to confirm the exp11 gain over gating.")
    doc.add_paragraph("  2. Integrate the token-level cross-attention model into the serving model in backend/models.")
    doc.add_paragraph("  3. Add inference-time target windowing (±40 nt around predicted binding site) for full-length mRNAs.")
    doc.add_paragraph("  4. Validate on an external dataset before deployment.")

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True) if os.path.dirname(OUTPUT_PATH) else None
    doc.save(OUTPUT_PATH)
    print(f"Word document saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
